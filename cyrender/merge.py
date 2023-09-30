# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import glob
import os
from decouple import config as env_config


def addEmptyParagraphs(doc, count):
    """
    Adds empty paragraphs in the number specified
    :param doc: The document object
    :param count: Number of empty paragraphs
    :return: The document object with the new empty paragraphs added
    """
    for i in range(0, count):
        doc.add_paragraph(" ")


def addElement(
    doc, para, lineSpacing, beforeSpacing, afterSpacing, indent, paraStyle, align
):
    """
    Adds a specified element to the document
    :param doc: The document object
    :param para: The text of the paragraph to add
    :param lineSpacing: The line spacing to use in the paragraph
    :param beforeSpacing: The extra space before the paragraph to use
    :param afterSpacing: The extra space after the paragraph to use
    :param indent: The amount to indent the text (e.g. first paragraphs in chapters should not be indented)
    :param paraStyle: The style to use for the paragraph. See here: https://python-docx.readthedocs.io/en/latest/user/styles-using.html
    :param align: The alignment to use for the paragraph. See here: https://python-docx.readthedocs.io/en/latest/api/text.html?highlight=alignment
    :return:
    """
    p = doc.add_paragraph(para, style=paraStyle)
    paragraphFormat = p.paragraph_format
    paragraphFormat.line_spacing_rule = lineSpacing
    paragraphFormat.space_before = Pt(beforeSpacing)
    paragraphFormat.space_after = Pt(afterSpacing)
    paragraphFormat.first_line_indent = Pt(indent)
    paragraphFormat.alignment = align


def processParagraph(doc, para: str, firstPara: bool):
    """
    Processes the paragraphs in the text
    :param doc: the document object
    :param para: the text of the paragraph being analysed
    :param firstPara: Bool value indicating if this is the first paragraph in the document or not
    :return: the document object
    """
    if para.strip() == "":
        return doc
    elif para.startswith("[MT]"):
        para = para.replace("[MT]", "")
        print("--> Adding book title: {}".format(para))
        addEmptyParagraphs(doc, 10)
        addElement(
            doc,
            para,
            WD_LINE_SPACING.SINGLE,
            0,
            0,
            0,
            "Title",
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    elif para.startswith("[By]"):
        para = para.replace("[By]", "By: ")
        print("--> Adding by line: {}".format(para))
        addElement(
            doc,
            para,
            WD_LINE_SPACING.SINGLE,
            0,
            0,
            0,
            "Subtitle",
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    elif para.startswith("[T]"):
        para = para.replace("[T]", "")
        print("--> Adding chapter title...{}".format(para))
        addEmptyParagraphs(doc, 4)
        addElement(
            doc,
            para,
            WD_LINE_SPACING.SINGLE,
            0,
            0,
            0,
            "Heading 1",
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    elif para.startswith("[S]"):
        para = para.replace("[S]", "")
        print("--> Adding chapter subtitle: {}".format(para))
        addElement(
            doc,
            para,
            WD_LINE_SPACING.DOUBLE,
            0,
            0,
            0,
            "Heading 2",
            WD_ALIGN_PARAGRAPH.CENTER,
        )
    elif para.strip() == "###":
        addEmptyParagraphs(doc, 1)
        print("--> Adding centered section break: {}...".format(para[0:50]))
        addElement(
            doc,
            para,
            WD_LINE_SPACING.DOUBLE,
            0,
            0,
            0,
            "Normal",
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        addEmptyParagraphs(doc, 1)
        firstPara = True
    else:
        print("--> Adding paragraph: {}...".format(para[0:50]))
        if not firstPara:
            addElement(
                doc,
                para,
                WD_LINE_SPACING.DOUBLE,
                0,
                0,
                18,
                "Normal",
                WD_ALIGN_PARAGRAPH.LEFT,
            )
        else:
            addElement(
                doc,
                para,
                WD_LINE_SPACING.DOUBLE,
                0,
                0,
                0,
                "Normal",
                WD_ALIGN_PARAGRAPH.LEFT,
            )
            firstPara = False
    return firstPara


def combineWordDocument(
    files: list, output_file: str, author_name: str, novel_title: str
):
    """
    Main controlling function that combines text files into a single word document
    :param files: A list of files to combine
    :param output_file: The output file (docx format) that the result will be saved as
    :param author_name: The name of the author
    :param novel_title: The title of the book
    :return: None. Saves the document and exits
    """
    mergedDocument = Document()
    paraFormat = mergedDocument.styles["Normal"].paragraph_format
    paraFormat.space_before = None

    for index, file in enumerate(files):
        print("Processing file {}...".format(file))
        if "chapter00" in file:
            section = mergedDocument.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            paragraph = header.paragraphs[0]
            paragraph.text = "{}\t{}".format(author_name, novel_title)
            paragraph.style = mergedDocument.styles["Header"]
        elif "chapter01" in file:
            section = mergedDocument.add_section()
            header = section.header
            header.is_linked_to_previous = False
            paragraph = header.paragraphs[0]
            paragraph.text = "{}\t{}\tPage #".format(author_name, novel_title)
            paragraph.style = mergedDocument.styles["Header"]
        filename = os.path.basename(file)
        fileTag = filename.replace("chapter", "").replace(".txt", "")
        with open(file, "r", encoding="utf-8") as myFile:
            myText = myFile.read()

        # Now, for each paragraph, add as a separate paragraph to new docx
        myParagraphs = myText.split("\n")
        firstPara = True

        for para in myParagraphs:
            if para.strip() == "":
                continue
            firstPara = processParagraph(mergedDocument, para, firstPara)

        # Don't add a page break if you've reached the last file.
        if index < len(files) - 1:
            mergedDocument.add_page_break()

    mergedDocument.save(output_file)


if __name__ == "__main__":
    file_pattern = env_config("FILE_PATTERN", default="./", cast=str)
    output_file = env_config("OUTPUT_FILE", cast=str)
    author_name = env_config("AUTHOR_NAME", cast=str)
    novel_title = env_config("TITLE", cast=str)
    files = sorted(glob.glob(file_pattern))
    combineWordDocument(files, output_file, author_name, novel_title)
